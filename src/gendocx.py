# -*- encoding: utf-8 -*- #
__author__ = 'FeiZhang 81597228@qq.com'
__date__ = '2019-07-20'

from docx import Document
import datetime


def gen_doc(title, author):
    """
    generator document structure
    :return:
    """
    currentDT = datetime.datetime.now()

    document = Document()

    document.add_heading(title, 0)
    # document.add_heading('The data table structure description', 0)

    document.add_paragraph('author:{}       date:{}'.format(author, currentDT.strftime("%Y-%m-%d")))
    # document.add_paragraph('')
    document.add_paragraph('\n')

    return document


def doc_append_table(doc, tb_info, tb_name, tb_comment):
    """
    append data table structure info to the document
    :param doc:
    :param tb_info:
    :param tb_name:
    :return:
    """
    table = doc.add_table(rows=1, cols=4, style='Table Grid')
    tbinfo_cells = table.rows[0].cells
    tbinfo_cells[0].merge(tbinfo_cells[1]).merge(tbinfo_cells[2]).merge(tbinfo_cells[3])
    tbinfo_cells[0].text = '表名:{}  {}'.format(tb_name, tb_comment)
    # tbinfo_cells = table.rows[1].cells
    # tbinfo_cells[0].merge(tbinfo_cells[1]).merge(tbinfo_cells[2]).merge(tbinfo_cells[3])
    # tbinfo_cells[0].text = '表描述:{}'.format(tb_name)
    # add table row
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = '列名'
    hdr_cells[1].text = '数据类型'
    # hdr_cells[2].text = 'Null'
    # hdr_cells[3].text = 'Key'
    hdr_cells[2].text = 'Default'
    hdr_cells[3].text = '栏位说明'

    for r in tb_info:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r[0])
        row_cells[1].text = str(r[1])
        # row_cells[2].text = str(r[3])
        # row_cells[3].text = str(r[4])
        row_cells[2].text = str(r[5])
        row_cells[3].text = str(r[8])

    # add space between table
    doc.add_paragraph("\n")


